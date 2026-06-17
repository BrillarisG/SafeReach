from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def touch_styles(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(10)


def update_backend_report() -> None:
    path = ROOT / "doc" / "backend" / "SafeReach_Backend_Implementation_Report.docx"
    document = Document(path)
    document.add_page_break()
    document.add_heading("DB-Backed Frontend, Sheet Export, and DevOps Phase", level=1)
    document.add_paragraph(f"Updated on {date.today().isoformat()}. This section records the latest implementation pass without exposing API keys, database passwords, tokens, or personal credentials.")

    document.add_heading("Completed Work", level=2)
    for item in [
        "Extended DB-1 with incident logs, safety reports, and API test result storage.",
        "Seeded one SafeReach school environment with stored classes, sections, teachers, students, timetable, attendance, reports, and incident data for frontend testing.",
        "Extended the backend bootstrap payload so frontend pages can read schools, classes, teachers, students, attendance, reports, incidents, API tests, and timetable data from DB-1.",
        "Added Google Sheets export for index tabs sheet-1, sheet-2, and sheet-3 plus separate per-table tabs such as sheet-1.1 schools and sheet-1.12 students.",
        "Added Dockerfiles, Docker Compose with Redis persistent volume, Render blueprint, Vercel configuration, Kubernetes manifests with two frontend/backend replicas, Redis PVC, Services, Ingress, and HPA.",
        "Added a combined GitHub Actions workflow for backend compile/smoke, frontend build, and Docker build checks on branch main.",
    ]:
        add_bullet(document, item)

    document.add_heading("API and Integration Test Results", level=2)
    add_table(
        document,
        ["Area", "Result", "Notes"],
        [
            ["DB-1 Supabase PostgreSQL", "Passed", "Connected and confirmed one school, six students, and two teachers."],
            ["DB-2 Supabase PostgreSQL", "Passed", "Verified protected DB-1 mirror tables and insert-only backup behavior."],
            ["MongoDB DB-3", "Passed", "Pinged and inserted system check event for realtime/event storage."],
            ["Google Sheets", "Passed", "Accessed workbook and exported 36 tabs using batch updates."],
            ["Twilio SMS", "Queued", "Test SMS was accepted by Twilio and queued for the configured test phone number."],
            ["Resend Email", "Blocked", "Resend rejected email because the sender/domain is not verified for the target recipient. Verify a domain or send to the allowed account email."],
            ["Redis", "Configured", "Supports REDIS_URL and Upstash Redis REST variables. Docker Compose and Kubernetes Redis configuration were added."],
        ],
    )

    document.add_heading("Frontend Stored Data Rule", level=2)
    for item in [
        "Admin dashboard, class records, staff directory, and timetable screens were updated to read stored backend bootstrap data instead of local hardcoded rows.",
        "When backend data is unavailable, updated screens show loading/error/empty states instead of silently displaying fake data.",
        "Write actions that do not yet have production backend APIs are marked as planned/read-only rather than creating local-only records that look stored.",
    ]:
        add_bullet(document, item)

    document.add_heading("Remaining Backend Work", level=2)
    for item in [
        "Expose authenticated WebSocket write events for teacher/student/admin mutations currently planned as read-only in the frontend.",
        "Add Redis-backed sessions, rate limiting, queues, and cache using REDIS_URL or Upstash Redis REST in each runtime.",
        "Verify a Resend sending domain and update RESEND_FROM_EMAIL before production email delivery.",
        "Create deployment secrets in Render, Vercel, GitHub Actions, and Kubernetes secret storage without committing them.",
    ]:
        add_bullet(document, item)
    touch_styles(document)
    document.save(path)


def update_frontend_report() -> None:
    path = ROOT / "doc" / "frontend" / "SafeReach_Frontend_Change_Report.docx"
    document = Document(path)
    document.add_page_break()
    document.add_heading("DB-Backed Display Phase", level=1)
    document.add_paragraph(f"Updated on {date.today().isoformat()}. This section records frontend changes that prevent unstored demo data from appearing as live records.")

    document.add_heading("Changed Screens", level=2)
    add_table(
        document,
        ["Screen", "Change"],
        [
            ["Admin Dashboard", "Reads schools, students, classes, incidents, reports, and attendance counts from backend bootstrap data."],
            ["Admin Class Records", "Displays stored classes, sections, and students only."],
            ["Admin Staff Directory", "Displays stored teacher and assignment records only; local add/edit actions are marked planned/read-only."],
            ["Timetable", "Loads timetable days, periods, and break positions from backend bootstrap data; no default timetable is shown when backend data is missing."],
        ],
    )

    document.add_heading("Verification Notes", level=2)
    for item in [
        "Google Sheet export now includes per-table stored data tabs and an API test summary tab.",
        "Frontend build should be run with NEXT_PUBLIC_SAFEREACH_API_URL pointing to the deployed backend API.",
        "Remaining pages that still need mutation APIs should avoid creating local-only data that appears persisted.",
    ]:
        add_bullet(document, item)
    touch_styles(document)
    document.save(path)


def main() -> None:
    update_backend_report()
    update_frontend_report()
    print("DOCX reports updated")


if __name__ == "__main__":
    main()
