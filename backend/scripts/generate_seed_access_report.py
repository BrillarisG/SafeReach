from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "doc" / "SafeReach_New_School_Dataset_and_Login_Report.docx"

TEACHERS = [
    ("Sarah Jenkins", "Primary incharge", "sarah.jenkins@safereach.school", "Class 4-A"),
    ("James Anderson", "Assistant incharge", "james.anderson@safereach.school", "Class 4-A"),
    ("Elena Roy", "Primary incharge", "elena.roy@safereach.school", "Class 5-A"),
    ("David Kumar", "Assistant incharge", "david.kumar@safereach.school", "Class 5-A"),
    ("Priya Nair", "Primary incharge", "priya.nair@safereach.school", "Class 6-A"),
    ("Ravi Menon", "Assistant incharge", "ravi.menon@safereach.school", "Class 6-A"),
    ("Anika Sharma", "Primary incharge", "anika.sharma@safereach.school", "Class 7-A"),
    ("Kiran Patel", "Assistant incharge", "kiran.patel@safereach.school", "Class 7-A"),
    ("Meera Iyer", "Primary incharge", "meera.iyer@safereach.school", "Class 8-A"),
    ("Arjun Singh", "Assistant incharge", "arjun.singh@safereach.school", "Class 8-A"),
]

PARENTS = [
    "Sarah Thompson", "Nisha Sharma", "Ravi Patel", "Meena Gupta", "Arun Kumar", "Kavya Iyer",
    "Sanjay Rao", "Lakshmi Nair", "Vikram Singh", "Divya Menon", "Rahul Verma", "Pooja Das",
    "Manoj Shah", "Anita Joshi", "Deepak Kapoor", "Renu Bhat", "Suresh Reddy", "Asha Pillai",
    "Naveen George", "Swathi Krishnan", "Harish Mehta", "Neha Malhotra", "Girish Jain", "Rekha Sethi",
    "Sunil Shetty", "Maya Chandra", "Rajesh Prasad", "Shilpa Bose", "Karthik Rao", "Anjali Thomas",
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(document: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    result = document.add_table(rows=1, cols=len(headers))
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    result.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = result.rows[0].cells[index]
        cell.width = Inches(widths[index])
        shade(cell, "0B2E6D")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
    for row in rows:
        cells = result.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            set_cell_text(cells[index], value)
    document.add_paragraph()


def heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(11, 46, 109)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(11, 46, 109)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SafeReach School Dataset and Access Report")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SafeReach Academy | DB-1 operational store and DB-2 immutable mirror")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(62, 90, 135)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(12)
    note_run = note.add_run("Test access only: replace every password before production. Do not store database, service, or API secrets in this document or in Git.")
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(176, 45, 35)

    heading(document, "Implementation Summary")
    document.add_paragraph("The new Supabase DB-1 and DB-2 projects were initialized with the same SafeReach schema. DB-1 is the editable operational store. DB-2 contains an insert-only mirror and blocks update and delete operations with database triggers.")
    table(document, ["Area", "Result"], [
        ("School", "SafeReach Academy (SAFE-ACADEMY-01)"),
        ("Users", "42 total: 1 main admin, 1 school administrator, 10 teachers, 30 parents"),
        ("Classes", "Class 4-A through Class 8-A"),
        ("Students", "100 stored student records; Leo and Maya Thompson share one parent account"),
        ("Attendance", "100 stored attendance records across all seeded students"),
        ("Timetable", "240 stored periods: 5 subjects, 8 periods, 6 days, and 3 movable breaks per class"),
        ("Backup", "DB-2 mirrored 646 inserted records; DB-2 update/delete guard passed"),
        ("Google Sheets", "54 index and per-table tabs synchronized"),
        ("JSON", "DB-1, DB-2, and DB-3 backup export completed"),
    ], [2.0, 4.9])

    heading(document, "Deployment Environment")
    document.add_paragraph("Render must use the new DB-1 and DB-2 connection strings in its Environment settings. Vercel must keep using the Render API URL only; the frontend must not connect directly to either database. The secret values are intentionally omitted from this report.")
    table(document, ["Service", "Required configuration"], [
        ("Render backend", "DB1_URL, DB2_URL, JWT_SECRET, SECRET_KEY, MongoDB, Redis, Google Sheets, Twilio, Resend, and Sentry values"),
        ("Vercel frontend", "VITE_SAFEREACH_API_URL and VITE_SAFEREACH_WS_URL pointing to the Render backend"),
        ("DB-1", "New Supabase project zcbxgqnmjpmsblccilpv; editable application data"),
        ("DB-2", "New Supabase project bgmelwpsktrisdjgixip; mirrored immutable backup data"),
    ], [1.65, 5.25])

    heading(document, "Role Login Access")
    document.add_paragraph("All accounts below are seeded test accounts. Each password is shared by role for testing only.")
    table(document, ["Role", "Name", "Email", "Test password"], [
        ("Main admin", "SafeReach Platform Owner", "main@safereach.local", "MainAdmin@2026"),
        ("School administrator", "Priya Raman", "admin@safereach.school", "SchoolAdmin@2026"),
    ], [1.35, 1.75, 2.75, 1.1])

    heading(document, "Teacher Accounts")
    table(document, ["Teacher", "Access", "Email", "Assignment", "Password"], [
        (name, access, email, assignment, "Teacher@2026") for name, access, email, assignment in TEACHERS
    ], [1.25, 1.25, 2.25, 1.2, 1.05])

    heading(document, "Parent Accounts")
    table(document, ["Parent", "Email", "Password"], [
        (name, f"parent{index:02d}@safereach.school", "Parent@2026") for index, name in enumerate(PARENTS, start=1)
    ], [2.4, 3.0, 1.6])

    heading(document, "Validation Result")
    document.add_paragraph("Passed: DB-1 connection, DB-2 connection, seeded count verification, DB-2 update/delete protection, Redis, Google Sheets, Twilio account validation, backend health, bootstrap, school-admin authentication, and JSON backup export.")
    document.add_paragraph("Needs operator action: MongoDB timed out from this machine, which usually means Atlas network access/DNS/TLS policy must be checked. Resend returned HTTP 401, which means the configured API key is no longer accepted and must be replaced in Render.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SafeReach - Internal test deployment handover")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(90, 100, 115)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
